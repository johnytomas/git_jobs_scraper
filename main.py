import requests
from bs4 import BeautifulSoup
import time
import os
from docx import Document
from docx.shared import Pt
import docx.oxml.shared
import docx.opc.constants

# Základní URL pro scraping - snadno upravitelné pro různé filtry
BASE_URL = "https://www.jobs.cz/prace/python-vyvojar/"

# Přidání headers pro requests
HEADERS = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36',
    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
    'Accept-Language': 'cs,en-US;q=0.7,en;q=0.3',
    'Connection': 'keep-alive',
}

def get_output_file_path():
    """Vrátí cestu k výstupnímu souboru"""
    current_directory = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(current_directory, "nabidky.docx")

def parse_job_listing(job_element):
    """Zpracuje jeden element nabídky práce a vrátí základní informace"""
    title = job_element.text.strip()
    link = job_element.find('a')['href']
    return title, link

def save_job_to_file(doc, title, company, address, salary, description, link):
    """Uloží informace o pracovní nabídce do Word dokumentu"""
    # Nadpis s větším písmem
    heading = doc.add_paragraph()
    run = heading.add_run(f"Pozice: {title}")
    run.font.size = Pt(20)
    run.bold = True
    
    # Základní informace
    doc.add_paragraph(f"Společnost: {company}")
    
    # Zpracování adres
    addresses = address.split('\n')
    p = doc.add_paragraph()
    p.add_run("Adresa: " + addresses[0])
    indent = " " * 8  # Zarovnání pod první adresu
    for addr in addresses[1:]:
        p.add_run('\n' + indent + addr)
    
    doc.add_paragraph(f"Plat: {salary}")
    
    # Detailní popis s lepším formátováním
    doc.add_paragraph()  # Prázdný řádek před detaily
    for line in description.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph()  # Prázdný řádek
        elif line.endswith(':'):  # Nadpisy a sekce
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
        elif line.startswith('- '):  # Odrážky
            doc.add_paragraph(line, style='List Bullet')
        else:  # Běžný text
            doc.add_paragraph(line)
    
    doc.add_paragraph()  # Prázdný řádek před odkazem
    
    # Přidání hypertextového odkazu
    p = doc.add_paragraph()
    p.add_run("Odkaz: ")
    add_hyperlink(p, link, link)
    
    # Přidání horizontální čáry
    paragraph = doc.add_paragraph()
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = docx.oxml.shared.OxmlElement('w:pBdr')
    bottom = docx.oxml.shared.OxmlElement('w:bottom')
    bottom.set(docx.oxml.shared.qn('w:val'), 'single')
    bottom.set(docx.oxml.shared.qn('w:sz'), '6')  # Tloušťka čáry (6 = 1/8 pt)
    bottom.set(docx.oxml.shared.qn('w:space'), '1')
    bottom.set(docx.oxml.shared.qn('w:color'), '000000')  # Černá barva
    pBdr.append(bottom)
    pPr.append(pBdr)
    
    # Nastavení mezer před a za čárou
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = docx.shared.Pt(12)
    paragraph_format.space_after = docx.shared.Pt(12)
    
    # Přidání prázdného řádku za oddělovačem
    doc.add_paragraph()

def add_hyperlink(paragraph, text, url):
    """Přidá hypertextový odkaz do odstavce"""
    # Tento kód přidá klikatelný odkaz do dokumentu
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Vytvoření hypertextového odkazu
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    
    # Vytvoření vlastností pro text odkazu
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')
    
    # Přidání modrého zbarvení
    c = docx.oxml.shared.OxmlElement('w:color')
    c.set(docx.oxml.shared.qn('w:val'), '0000FF')
    rPr.append(c)
    
    # Přidání podtržení
    u = docx.oxml.shared.OxmlElement('w:u')
    u.set(docx.oxml.shared.qn('w:val'), 'single')
    rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    
    return hyperlink

def parse_job_description(details_job):
    """Zpracuje detailní popis pracovní pozice"""
    if not details_job:
        return "Detaily pozice: Neuvedeno"
    
    job_description = ""
    current_section = None
    elements = details_job.find_all(['h2', 'h3', 'p', 'ul', 'li', 'div', 'strong'], recursive=True)
    
    for element in elements:
        # Zpracování nadpisů
        if element.name in ['h2', 'h3'] or 'heading' in element.get('class', []):
            current_section = element.text.strip(':').strip()
            job_description += f"\n{current_section}:\n"
        
        # Zpracování odstavců
        elif element.name == 'p':
            text = element.text.strip()
            if text:
                # Kontrola, zda odstavec obsahuje tučný text
                if element.find('strong') or element.find('b'):
                    job_description += f"\n{text}\n"
                else:
                    job_description += f"{text}\n"
        
        # Zpracování seznamů
        elif element.name == 'ul':
            job_description += "\n"
            for li in element.find_all('li', recursive=False):
                job_description += f"- {li.text.strip()}\n"
            job_description += "\n"
        
        # Zpracování samostatných odrážek
        elif element.name == 'li' and not element.find_parent('ul', recursive=False):
            job_description += f"- {element.text.strip()}\n"
        
        # Zpracování tučného textu
        elif element.name == 'strong' or element.name == 'b':
            if not element.find_parent(['h2', 'h3', 'p']):
                job_description += f"\n{element.text.strip()}:\n"
        
        # Zpracování speciálních div elementů
        elif element.name == 'div' and 'section' in element.get('class', []):
            text = element.text.strip()
            if text:
                job_description += f"\n{text}\n"
    
    # Vyčištění vícenásobných prázdných řádků
    job_description = '\n'.join(line for line in job_description.splitlines() if line.strip())
    
    return job_description.strip()

def format_addresses(addresses):
    """Formátuje seznam adres se zarovnáním pod první adresu"""
    if not addresses:
        return "Neuvedeno"
    
    address_list = list(set(addr.text.strip() for addr in addresses))
    address_list.sort()
    
    # První adresa začíná normálně, další jsou odsazené
    if len(address_list) == 1:
        return address_list[0]
    else:
        # Vytvoříme první řádek s "Adresa: "
        result = f"Adresa: {address_list[0]}\n"
        # Přidáme další adresy se stejným odsazením jako má první adresa
        padding = " " * 8  # 8 mezer odpovídá délce "Adresa: "
        result += "\n".join(f"{padding}{addr}" for addr in address_list[1:])
        return result

def get_job_details(url):
    """Získá detailní informace z jednotlivého inzerátu"""
    try:
        response = requests.get(url, headers=HEADERS)  # Přidání headers
        soup = BeautifulSoup(response.text, 'html.parser')
        
        company_element = soup.find('div', class_="IconWithText")
        name_company = company_element.find('p', class_="typography-body-medium-text-regular").text.strip() if company_element else "Neuvedeno"
        
        address = format_addresses(soup.find_all('a', class_="link-secondary link-underlined"))
        
        salary_element = soup.find("div", {"data-test": "jd-salary"})
        salary = salary_element.find("p", class_="typography-body-medium-text-regular").text.strip() if salary_element else "Neuvedeno"
        
        details_element = soup.find("div", {"data-test": "jd-body-richtext"}, class_="RichContent mb-1400")
        job_description = parse_job_description(details_element)
        
        return name_company, address, salary, job_description
        
    except Exception as e:
        print(f"Chyba při získávání detailů: {e}")
        return "Neuvedeno", "Neuvedeno", "Neuvedeno", "Detaily pozice: Neuvedeno"

def scrape_jobs():
    """Hlavní funkce pro stahování pracovních nabídek"""
    total_jobs = 0
    current_url = BASE_URL
    page = 1
    
    doc = Document()
    
    while True:
        try:
            print(f"Stahuji stránku {page}...")
            response = requests.get(current_url, headers=HEADERS)
            soup = BeautifulSoup(response.text, 'html.parser')
            
            jobs = soup.find_all('h2', class_="SearchResultCard__title")
            if not jobs:
                print("Nenašli jsme žádné nabídky práce odpovídající zadání.")
                break

            for job in jobs:
                title, link = parse_job_listing(job)
                print(f"Stahuji detail inzerátu: {title}")
                
                name_company, address, salary, job_description = get_job_details(link)
                save_job_to_file(doc, title, name_company, address, salary, job_description, link)
                time.sleep(0.25)
            
            total_jobs += len(jobs)
            print(f"Staženo {len(jobs)} nabídek ze stránky {page}")

            next_page = soup.find('a', class_='Button Button--secondary Button--square Pagination__button--next')
            
            if not next_page:
                print("Dosaženo poslední stránky")
                break
            
            current_url = next_page['href']
            if not current_url.startswith('http'):
                current_url = 'https://www.jobs.cz' + current_url
            
            page += 1
            
        except Exception as e:
            print(f"Nastala chyba při stahování: {e}")
            break
    
    doc.save(get_output_file_path())
    print(f"\nCelkem nalezeno {total_jobs} pracovních nabídek")
    print("Výsledky byly uloženy do souboru 'nabidky.docx'")

if __name__ == '__main__':
    scrape_jobs()
